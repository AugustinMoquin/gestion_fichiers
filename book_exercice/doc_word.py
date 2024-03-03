from docx import Document
from docx.shared import Inches
from PIL import Image
from collections import Counter

def create_word_document(docx_filename, image_paths, title, auth, word_count_list):
    doc = Document()
    doc.add_heading("Rendu Projet Python", level=1)
    
    doc.add_heading(f"Title : {title}", level=3)
    doc.add_heading(f"Author : {auth}", level=3)

    for image_path in image_paths:
        image = Image.open(image_path)
        max_width = Inches(5)  # Adjust the maximum width as needed
        if image.width > max_width:
            ratio = max_width / image.width
            new_height = ratio * image.height
            image = image.resize((max_width, new_height))
            
    doc.add_paragraph("image origale")
    doc.add_picture(image_paths[0], width=max_width)
    
    doc.add_paragraph("image modifiée")
    doc.add_picture(image_paths[1], width=max_width)

    doc.add_paragraph("Fait par : augustin moquin")
    
    doc.add_page_break()
    
    doc.add_picture("./img/word_count_plot.png", width=max_width)
    doc.add_paragraph(f'In this plot we can see that out of the {len(word_count_list)}, the most frequent number of words is {most_frequent_number(word_count_list)}.\n We find that the min number of word in a paragraph is {min(word_count_list)}, and the max is {max(word_count_list)} ')
    
    doc.save(docx_filename)
    
def most_frequent_number(lst):
    count_dict = Counter(lst)
    most_common = count_dict.most_common(1)
    return most_common[0][0] if most_common else None

def doc_main(title, auth, word_count_list):
    # Example usage
    output_file = "output.docx"
    images = ["./img/Henry_Jones.jpg", "./img/resized_Henry_Jones.jpg"] 
    

    create_word_document(output_file, images, title, auth, word_count_list)