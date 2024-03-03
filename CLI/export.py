import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import customtkinter
from global_var import *
from db_connect import *

conn = db_connect()
cursor = conn.cursor()

def fetch_data(table_name):
    try :
        #--------------------------------#
        query = f"SELECT * FROM {table_name}"
        cursor.execute(query)

        rows = cursor.fetchall()
        #--------------------------------#
        query = f"PRAGMA table_info({table_name})"
        cursor.execute(query)

        columns_info = cursor.fetchall()

        column_names = [column[1] for column in columns_info]
    except Exception as e:
        print(f"there was an error : {e}")
    return rows, column_names



def combobox_callback_column1(choice):
    #rend la var global pour la récuperer et le comparer avec le choice 2
    global choice1
    choice1 = choice
    # try :
    if choice != None and choice2 != None:
        graph_creation(choice, choice2)
    else:
        pass
    # except Exception:
    #     pass

def combobox_callback_column2(choice):
    global choice2
    choice2 = choice
    # try:
    if choice != None and choice1 != None:
        graph_creation(choice, choice1)

    else:
        pass
    # except Exception:
    #     pass
    

def combobox_callback_choice(choice):
    #get the table column to choose
    _rows, _columns = fetch_data(choice)
    global rows
    global columns
    global combobox_column1
    global combobox_column2
    global column1_label
    global column2_label
    columns = _columns
    rows = _rows
    
    column1_label = customtkinter.CTkLabel(Vwindow, text="choose a column", font=custom_font)
    column1_label.pack(pady=10)
    combobox_column1 = customtkinter.CTkComboBox(master=Vwindow,
                                        values=columns,
                                        command=combobox_callback_column1)
    combobox_column1.pack(padx=20, pady=10)
    combobox_column1.set("")
    
    column2_label = customtkinter.CTkLabel(Vwindow, text="choose a column", font=custom_font)
    column2_label.pack(pady=10)
    combobox_column2 = customtkinter.CTkComboBox(master=Vwindow,
                                        values=columns,
                                        command=combobox_callback_column2,
                                        font=custom_font)
    combobox_column2.pack(padx=20, pady=10)
    combobox_column2.set("")
    

def param_choice(window, table_name):
    global Vwindow
    Vwindow = window
    choice_label = customtkinter.CTkLabel(window, text="choose a table", font=custom_font)
    choice_label.pack(pady=10)
    combobox_choice = customtkinter.CTkComboBox(master=window,
                                        values=table_name,
                                        command=combobox_callback_choice,
                                        font=custom_font)
    combobox_choice.pack(padx=20, pady=10)
    combobox_choice.set("")

def graph_creation(choice1, choice2):
    for i, v in enumerate(columns):
        if choice1 == v:
            x = [row[i] for row in rows]
        if choice2 == v:
            y = [row[i] for row in rows]
            
    plt.plot(x, y)
    plt.xlabel(choice1)
    plt.ylabel(choice2)
    plt.title('plot export')

    # Save the plot as an image file
    image_path = 'matplotlib_plot.png'
    plt.savefig(image_path, format='png')
    
    # Clear the plot for reuse
    plt.clf()
    #creer un nv doc word
    doc = Document()
    
    graph_export(doc, image_path)
    doc.save('output_with_matplotlib_plot.docx')
    #delete the old combobox
    combobox_column1.pack_forget()
    combobox_column2.pack_forget()
    column1_label.pack_forget()
    column2_label.pack_forget()

    label = customtkinter.CTkLabel(Vwindow, text='generation complete')
    
    label.pack()

def graph_export(document, image_path):
    document.add_paragraph('Matplotlib Plot:')
    document.add_picture(image_path, width=Inches(4))
    document.add_paragraph('Description or caption for the plot.')